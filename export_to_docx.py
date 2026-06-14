"""
Export reboot-camp index.html → reboot-camp-edit.docx
Clean Word doc for author review and redlining.
"""

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, re

SRC = os.path.join(os.path.dirname(__file__), "src", "index.html")
OUT = os.path.join(os.path.dirname(__file__), "output", "reboot-camp-edit.docx")

os.makedirs(os.path.dirname(OUT), exist_ok=True)

with open(SRC, encoding="utf-8") as f:
    soup = BeautifulSoup(f.read(), "lxml")

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────

def style_normal(p, size=11, color=None, bold=False, italic=False):
    run = p.runs[0] if p.runs else p.add_run()
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_spacing(p, before=0, after=8, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        from docx.shared import Pt as DPt
        pf.line_spacing = DPt(line)

def add_heading(doc, text, level, size, bold=True, color=(15, 23, 42)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    set_spacing(p, before=18 if level == 1 else 12, after=6)
    return p

def add_body(doc, text):
    if not text.strip():
        return None
    p = doc.add_paragraph(text.strip())
    run = p.runs[0] if p.runs else p.add_run()
    run.font.size = Pt(11)
    set_spacing(p, before=0, after=6)
    return p

def add_pull_quote(doc, text, cite=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    set_spacing(p, before=10, after=4)
    run = p.add_run(f'"{text.strip()}"')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(6, 182, 212)  # cyan
    if cite:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.4)
        set_spacing(p2, before=0, after=10)
        r = p2.add_run(cite.strip())
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = RGBColor(71, 85, 105)  # slate

def add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    set_spacing(p, before=8, after=8)
    run = p.add_run(f"▌  {text.strip()}")
    run.font.size = Pt(11)
    run.font.bold = True

def add_affirmation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=10, after=10)
    run = p.add_run(text.strip())
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.bold = True

def add_statement(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=10, after=10)
    run = p.add_run(text.strip())
    run.font.size = Pt(13)
    run.font.bold = True

def add_action(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=12, after=4)
    label = p.add_run("THE ACTION  ")
    label.font.size = Pt(9)
    label.font.bold = True
    label.font.color.rgb = RGBColor(6, 182, 212)
    body = p.add_run(text.strip())
    body.font.size = Pt(11)

def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=10, after=10)
    run = p.add_run(text.strip())
    run.font.size = Pt(11)
    run.font.bold = True

def add_separator(doc):
    p = doc.add_paragraph("─" * 60)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    p.runs[0].font.color.rgb = RGBColor(200, 200, 200)

def get_inline_text(el):
    """Render inline HTML preserving bold/italic as plain text marks."""
    parts = []
    for child in el.children:
        if isinstance(child, NavigableString):
            parts.append(str(child))
        elif child.name in ("strong", "b"):
            parts.append(f"**{child.get_text()}**")
        elif child.name in ("em", "i"):
            parts.append(f"_{child.get_text()}_")
        elif child.name == "br":
            parts.append("\n")
        else:
            parts.append(child.get_text())
    return "".join(parts)

def add_para_with_inline(doc, el):
    text = get_inline_text(el).strip()
    if not text:
        return
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=6)
    # simple bold/italic rendering
    segments = re.split(r'(\*\*.*?\*\*|_.*?_)', text)
    for seg in segments:
        if seg.startswith("**") and seg.endswith("**"):
            run = p.add_run(seg[2:-2])
            run.font.bold = True
            run.font.size = Pt(11)
        elif seg.startswith("_") and seg.endswith("_"):
            run = p.add_run(seg[1:-1])
            run.font.italic = True
            run.font.size = Pt(11)
        else:
            run = p.add_run(seg)
            run.font.size = Pt(11)

# ── Process each section / article ────────────────────────────────────────────

def process_body_text(doc, div):
    for child in div.children:
        if isinstance(child, NavigableString):
            continue
        tag = child.name
        cls = " ".join(child.get("class", []))

        if tag == "p":
            add_para_with_inline(doc, child)
        elif tag in ("ul", "ol"):
            for li in child.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.3)
                set_spacing(p, before=0, after=4)
                run = p.add_run(get_inline_text(li).strip())
                run.font.size = Pt(11)
        elif tag == "blockquote" and "pull-quote" in cls:
            qt = child.find("p")
            cite = child.find("cite")
            add_pull_quote(doc, qt.get_text() if qt else child.get_text(),
                           cite.get_text() if cite else None)
        elif "callout" in cls:
            add_callout(doc, child.get_text(separator=" ", strip=True))
        elif "affirmation" in cls:
            add_affirmation(doc, child.get_text(separator=" ", strip=True))
        elif "statement" in cls:
            add_statement(doc, child.get_text(separator=" ", strip=True))
        elif "the-action" in cls:
            action_p = child.find("p", class_="action-text")
            if action_p:
                add_action(doc, action_p.get_text(strip=True))
        elif "formula" in cls:
            add_formula(doc, child.get_text(strip=True))
        elif "stat-block" in cls:
            items = child.find_all(class_="stat-item")
            row = "  ·  ".join(
                f"{i.find(class_='stat-number').get_text()} {i.find(class_='stat-label').get_text()}"
                for i in items if i.find(class_="stat-number")
            )
            add_formula(doc, row)
        elif "sign-off" in cls:
            p = doc.add_paragraph()
            set_spacing(p, before=12, after=6)
            run = p.add_run(child.get_text(separator=" ", strip=True))
            run.font.italic = True
            run.font.size = Pt(11)
        elif tag == "hr":
            add_separator(doc)
        elif tag in ("div", "section", "article"):
            process_body_text(doc, child)


# ── Cover ─────────────────────────────────────────────────────────────────────
cover = soup.find("section", id="cover")
if cover:
    add_heading(doc, "REBOOT CAMP", 1, 28)
    subtitle_el = cover.find("h1", class_="cover-title")
    if subtitle_el:
        add_heading(doc, subtitle_el.get_text(separator=" ", strip=True), 2, 16, bold=False)
    author_el = cover.find(class_="cover-overline")
    if author_el:
        p = doc.add_paragraph(author_el.get_text(strip=True))
        p.runs[0].font.size = Pt(12)
        p.runs[0].font.color.rgb = RGBColor(71, 85, 105)
    doc.add_page_break()

# ── Sections & Chapters (in document order) ──────────────────────────────────
# Iterate body children skipping nav, cover, toc, footer, script
body = soup.body
skip_ids = {"cover", "toc"}
skip_tags = {"nav", "footer", "script"}

for el in body.children:
    if isinstance(el, NavigableString):
        continue
    if el.name in skip_tags:
        continue
    el_id = el.get("id", "")
    el_cls = " ".join(el.get("class", []))
    if el_id in skip_ids:
        continue

    # progress bar, overlay divs — skip
    if el_id in ("progress-bar", "nav-overlay"):
        continue

    # image dividers between chapters — skip
    if "image-wrap" in el_cls and el.name == "div":
        continue

    # ── Section openers (note, how-to, origin, step openers, closing) ──
    if el.name == "section" and ("section-opener" in el_cls or "chapter-opener" in el_cls):
        add_separator(doc)
        step_num = el.find(class_="step-number")
        step_label = el.find(class_="step-label")
        h = el.find(["h2"])
        title_text = h.get_text(strip=True) if h else ""

        if step_num and step_label:
            p = doc.add_paragraph()
            set_spacing(p, before=6, after=2)
            run = p.add_run(f"STEP {step_num.get_text(strip=True)}  ·  {step_label.get_text(strip=True).upper()}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(6, 182, 212)

        if title_text:
            add_heading(doc, title_text, 2, 20)

        # pull quote inside opener
        pq = el.find("blockquote", class_="pull-quote")
        if pq:
            qt = pq.find("p")
            cite = pq.find("cite")
            add_pull_quote(doc, qt.get_text() if qt else pq.get_text(),
                           cite.get_text() if cite else None)

        # closing section (no id, has accent-rule + body-text)
        if not el_id:
            body_div = el.find(class_="body-text")
            if body_div:
                process_body_text(doc, body_div)
            sign_off = el.find(class_="sign-off")
            if sign_off:
                p = doc.add_paragraph()
                set_spacing(p, before=12, after=6)
                run = p.add_run(sign_off.get_text(separator=" ", strip=True))
                run.font.italic = True

        # section-openers with no step num (note, how-to, origin)
        section_titles = {"note-title", "howto-title", "origin-title"}
        if h and h.get("id") in section_titles:
            pass  # title already added above

        continue

    # ── Articles (chapter body content) ──
    if el.name == "article" and "chapter" in el_cls:
        ebook_body = el.find(class_="ebook-body")
        if ebook_body:
            process_body_text(doc, ebook_body)
        continue

    # ── About section ──
    if el.name == "section" and el_id == "about":
        add_separator(doc)
        add_heading(doc, "About the Author", 2, 18)
        name = el.find(class_="author-name")
        role = el.find(class_="author-role")
        if name:
            p = doc.add_paragraph(name.get_text(strip=True))
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.bold = True
        if role:
            p = doc.add_paragraph(role.get_text(strip=True))
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.color.rgb = RGBColor(71, 85, 105)
        about_text = el.find(class_="about-text")
        if about_text:
            for p_el in about_text.find_all("p"):
                add_body(doc, p_el.get_text(strip=True))
        continue

doc.save(OUT)
print(f"Saved: {OUT}")
